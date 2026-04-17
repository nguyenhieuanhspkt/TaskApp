# -*- coding: utf-8 -*-
import os
import platform
import ctypes
import torch
import pandas as pd
import re
import time
import tempfile
import hashlib
import pickle
from importlib.util import find_spec
from docx import Document
from PyQt5.QtCore import QThread, pyqtSignal
from rapidfuzz import fuzz
from sentence_transformers import SentenceTransformer, util

# ✅ NEW: import bộ làm sạch
from datacleaner import (
    clean_master_df,
    clean_items,
    normalize_text,  # dùng thống nhất cho SmartMatcher
)

# --- 1) FIX LỖI DLL WINERROR 1114 (Windows) ---
if platform.system() == "Windows":
    try:
        spec = find_spec("torch")
        if spec and spec.origin:
            dll_path = os.path.join(os.path.dirname(spec.origin), "lib", "c10.dll")
            if os.path.exists(dll_path):
                ctypes.CDLL(os.path.normpath(dll_path))
    except:
        pass


# --- 2) BỘ MÁY SO KHỚP THÔNG MINH (BGE-M3 HYBRID AI) ---
class SmartMatcher:
    def __init__(self, master_list, master_path, log_callback, model_path):
        """
        master_list: list[{"Ma","Ten","TSKT","DVT","Full_Norm"}]
        master_path: đường dẫn file Excel để hash cache
        model_path : đường dẫn model BGE-M3 (local) hoặc tên model
        """
        self.master_list = master_list
        self.texts = [m["Full_Norm"] for m in master_list]

        log_callback(f"🤖 Đang nạp Bộ não AI từ: {os.path.basename(str(model_path))}...")

        # Load model (CPU)
        self.model = SentenceTransformer(model_path, device="cpu")

        # --- Cơ chế caching embeddings ---
        cache_dir = os.path.join(tempfile.gettempdir(), "audit_cache_v2")
        os.makedirs(cache_dir, exist_ok=True)

        with open(master_path, "rb") as f:
            file_hash = hashlib.md5(f.read()).hexdigest()

        self.cache_path = os.path.join(cache_dir, f"bge_m3_{file_hash}.pkl")

        if os.path.exists(self.cache_path):
            log_callback("🚀 Đã tìm thấy bộ nhớ đệm AI! Nạp dữ liệu trong 1 giây...")
            with open(self.cache_path, "rb") as f:
                self.embeddings = pickle.load(f)
        else:
            log_callback("🧠 AI đang học dữ liệu Kho lần đầu (CPU mất 5-15 phút, vui lòng đợi)...")
            self.embeddings = self.model.encode(
                self.texts, convert_to_tensor=True, show_progress_bar=False
            )
            with open(self.cache_path, "wb") as f:
                pickle.dump(self.embeddings, f)
            log_callback("✅ Đã lưu bộ nhớ đệm cho các lần chạy sau.")

    def match(self, query: str):
        """
        Trả về (bản_ghi_tốt_nhất | None, điểm 0-100)
        Sử dụng hybrid: cosine semantic + fuzzy + penalty từ khoá hoá chất.
        """
        query_norm = normalize_text(query)
        if not query_norm:
            return None, 0

        query_vec = self.model.encode(query_norm, convert_to_tensor=True)
        cos_scores = util.cos_sim(query_vec, self.embeddings)[0]
        top_k = min(5, len(self.master_list))
        top_results = torch.topk(cos_scores, k=top_k)

        best_item, best_score = None, 0
        keywords = ["axit", "acid", "javen", "clo", "soda", "hcl", "clo2", "phenolphthalein"]

        for score_tensor, idx_tensor in zip(top_results[0], top_results[1]):
            idx = idx_tensor.item()
            sem_score = score_tensor.item() * 100
            m = self.master_list[idx]
            target_norm = m["Full_Norm"]

            fuzzy_score = fuzz.token_set_ratio(query_norm, target_norm)

            # penalty nếu keyword hoá chất xuất hiện lệch giữa query & target
            penalty = 0
            for kw in keywords:
                if (kw in query_norm) != (kw in target_norm):
                    penalty += 40

            final_score = (sem_score * 0.6) + (fuzzy_score * 0.4) - penalty

            if final_score > best_score:
                best_score = final_score
                best_item = m

        if best_score < 50:
            return None, int(best_score)
        return best_item, int(best_score)


# --- 3) LUỒNG XỬ LÝ CHÍNH ---
class WorkerThread(QThread):
    progress = pyqtSignal(int)
    log_signal = pyqtSignal(str)
    result_ready = pyqtSignal(object)

    def __init__(self, master_path, docx_path, model_path):
        super().__init__()
        self.master_path = master_path
        self.docx_path = docx_path
        self.model_path = model_path  # đường dẫn model từ GUI

    def run(self):
        word_app = None
        temp_kho_path = ""
        try:
            # 3.1) Đọc & làm sạch danh mục Excel
            self.log_signal.emit("📊 Đang đọc danh mục Excel...")
            # Gợi ý engine='openpyxl' cho .xlsx để ổn định
            df_master_raw = pd.read_excel(self.master_path, engine="openpyxl").fillna("")

            self.log_signal.emit("🧽 Đang làm sạch danh mục...")
            df_master_clean = clean_master_df(
                df_master_raw,
                log=self.log_signal.emit,
                fuzzy_dedup=True,
                fuzzy_threshold=92,
            )

            # chuyển sang dạng list dict cho SmartMatcher
            master_list = []
            for _, row in df_master_clean.iterrows():
                master_list.append(
                    {
                        "Ma": str(row["Ma"]),
                        "Ten": str(row["Ten"]),
                        "TSKT": str(row["TSKT"]),
                        "DVT": str(row["DVT"]),
                        "Full_Norm": str(row["Full_Norm"]),
                    }
                )

            # 3.2) Khởi tạo matcher
            matcher = SmartMatcher(
                master_list, self.master_path, self.log_signal.emit, self.model_path
            )

            # 3.3) Trích & làm sạch dữ liệu từ Word
            self.log_signal.emit("📄 Đang phân tích file Hồ sơ...")
            doc_hoso = Document(self.docx_path)
            raw_items = []
            for table in doc_hoso.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    # ràng buộc: có STT dạng số đầu dòng + có tên đủ dài
                    if len(cells) >= 3 and re.match(r"^\d+", cells[0]) and len(cells[1]) > 2:
                        raw_items.append(
                            {
                                "stt": cells[0],
                                "ten": cells[1],
                                "tskt": cells[2],
                                "dvt": cells[3] if len(cells) > 3 else "",
                            }
                        )

            if not raw_items:
                raise Exception("Không tìm thấy bảng vật tư phù hợp trong Word!")

            self.log_signal.emit("🧽 Đang làm sạch dữ liệu từ hồ sơ...")
            items = clean_items(raw_items, log=self.log_signal.emit)

            # 3.4) Tạo file tạm “Kho Chuẩn” để so sánh Track Changes
            from docx.shared import Pt

            doc_kho = Document()
            table_kho = doc_kho.add_table(rows=1, cols=4)
            table_kho.style = "Table Grid"
            headers = ["STT", "Tên vật tư", "Thông số kỹ thuật", "ĐVT"]
            for i, h in enumerate(headers):
                table_kho.rows[0].cells[i].text = h

            # 3.5) Matching
            results = []
            n = len(items)
            for i, item in enumerate(items):
                if self.isInterruptionRequested():
                    return

                self.log_signal.emit(f"🔍 Đang khớp mục {item['stt']}: {item['ten'][:30]}...")
                query = f"{item['ten']} {item['tskt']}".strip()

                best_match, score = matcher.match(query)

                row_cells = table_kho.add_row().cells
                row_cells[0].text = item["stt"]

                if best_match and score >= 55:
                    row_cells[1].text = best_match["Ten"]
                    row_cells[2].text = best_match["TSKT"]
                    row_cells[3].text = best_match["DVT"]
                    results.append(
                        {
                            "STT": item["stt"],
                            "Trạng thái": "✅ Khớp",
                            "Mã": best_match["Ma"],
                            "Điểm": f"{score}%",
                        }
                    )
                else:
                    row_cells[1].text = "KHÔNG TÌM THẤY TRONG KHO CHUẨN"
                    row_cells[2].text = item["tskt"]
                    row_cells[3].text = item["dvt"]
                    results.append(
                        {
                            "STT": item["stt"],
                            "Trạng thái": "❌ Lệch",
                            "Mã": "N/A",
                            "Điểm": f"{score}%",
                        }
                    )

                self.progress.emit(int((i + 1) / max(n, 1) * 85))

            # 3.6) Lưu tạm & So sánh Track Changes bằng Word COM (Windows)
            docx_dir = os.path.dirname(os.path.abspath(self.docx_path))
            output_path = os.path.join(docx_dir, f"KetQua_ThamDinh_{int(time.time())}.docx")
            temp_kho_path = os.path.join(tempfile.gettempdir(), f"tkho_{int(time.time())}.docx")
            doc_kho.save(temp_kho_path)

            self.log_signal.emit("📊 Đang so sánh Track Changes...")
            if platform.system() != "Windows":
                # Không có Word COM trên non-Windows: chỉ lưu “Kho chuẩn” để người dùng tự so sánh
                self.log_signal.emit("ℹ️ Môi trường không phải Windows. Bỏ qua Word COM compare.")
                # Gợi ý: Sao chép file tạm sang output_path
                import shutil

                shutil.copyfile(temp_kho_path, output_path)
            else:
                try:
                    import win32com.client

                    try:
                        word_app = win32com.client.GetActiveObject("Word.Application")
                    except:
                        word_app = win32com.client.Dispatch("Word.Application")

                    word_app.Visible = False
                    d_orig = word_app.Documents.Open(os.path.abspath(temp_kho_path))
                    d_rev = word_app.Documents.Open(os.path.abspath(self.docx_path))

                    d_comp = word_app.CompareDocuments(
                        OriginalDocument=d_orig, RevisedDocument=d_rev, Destination=2, Granularity=1
                    )
                    d_comp.SaveAs2(output_path)

                    d_orig.Close(0)
                    d_rev.Close(0)
                    d_comp.Close(-1)

                    # Mở file kết quả
                    try:
                        os.startfile(output_path)  # type: ignore[attr-defined]
                    except Exception:
                        pass
                except Exception as e:
                    self.log_signal.emit(f"⚠️ Không thể thực hiện Word COM compare: {e}")
                    # fallback: lưu file tạm thành output
                    import shutil

                    shutil.copyfile(temp_kho_path, output_path)

            self.progress.emit(100)
            self.result_ready.emit(pd.DataFrame(results))
            self.log_signal.emit("✅ Xong! File đã lưu tại thư mục hồ sơ.")

        except Exception as e:
            self.log_signal.emit(f"❌ Lỗi: {str(e)}")
        finally:
            if platform.system() == "Windows":
                try:
                    # nếu đã khởi tạo word_app thì quit
                    if "word_app" in locals() and word_app:
                        word_app.Quit()
                except:
                    pass
            if temp_kho_path and os.path.exists(temp_kho_path):
                try:
                    os.remove(temp_kho_path)
                except:
                    pass