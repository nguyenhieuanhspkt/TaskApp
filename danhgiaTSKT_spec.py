import re
import json
import os
import unicodedata
import time
import requests
import pandas as pd
from docx import Document
from bs4 import BeautifulSoup
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Tuple

# ======================
# 1) Cấu hình nhận diện
# ======================
@dataclass
class SpecsConfig:
    model_col_keys: set = field(default_factory=lambda: {"model", "pn", "part number", "so hieu", "ma san pham", "ma may"})
    brand_col_keys: set = field(default_factory=lambda: {"hang", "brand", "make", "thuong hieu", "hang san xuat"})
    mfr_col_keys: set = field(default_factory=lambda: {"nha san xuat", "nha sx", "manufacturer", "mfr"})
    
    # Cột mặc định để quét văn bản
    default_text_cols: List[str] = field(default_factory=lambda: [
        "Thông số kỹ thuật", "Mô tả", "Description", "Specs", "Danh mục hàng hóa", "Tên hàng"
    ])
    
    brand_blacklist: set = field(default_factory=lambda: {
        "chen lam kin", "gioang", "phu tung", "thiet bi", "vat tu", "may moc", "hang hoa"
    })
    
    unit_blacklist: set = field(default_factory=lambda: {
        "mm", "cm", "inch", "lbs", "kg", "watt", "volt", "hz", "amp", "bar", "psi"
    })
    
    extract_patterns: Dict = field(default_factory=lambda: {
        "model_labels": [r"(?:\bmodel\b|\bpn\b|\bpart number\b)\s*[:\-]?\s*([A-Za-z0-9\-\._\/]+)"],
        "model_freeform": r"\b([A-Z]{1,}[A-Z0-9\-]{2,}|[0-9]{2,}[A-Z]{1,}[A-Z0-9\-]*)\b"
    })

# ======================
# 2) Các hàm bổ trợ
# ======================
def load_cache() -> Dict:
    if os.path.exists("brand_cache.json"):
        try:
            with open("brand_cache.json", "r", encoding="utf-8") as f:
                return json.load(f)
        except: return {}
    return {}

def save_cache(cache_data: Dict):
    with open("brand_cache.json", "w", encoding="utf-8") as f:
        json.dump(cache_data, f, ensure_ascii=False, indent=4)

def norm_key(s: str) -> str:
    if not s: return ""
    s = str(s).lower().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    return " ".join(s.split())

def non_empty(x) -> bool:
    return x is not None and str(x).strip() != "" and not (isinstance(x, float) and pd.isna(x))

def fetch_brand_from_dauthau(model_str: str) -> Tuple[Optional[str], int]:
    if not model_str or len(model_str) < 3: return None, 0
    url = f"https://dauthau.asia/hanghoa/?key={model_str}"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        time.sleep(0.8)
        res = requests.get(url, headers=headers, timeout=10)
        if res.status_code == 200:
            soup = BeautifulSoup(res.content, 'html.parser')
            rows = soup.find_all('tr')
            count = len(rows) - 1 if len(rows) > 0 else 0
            if len(rows) >= 2:
                cols = rows[1].find_all('td')
                if len(cols) >= 3:
                    return cols[2].get_text(strip=True), count
            return None, count
    except: pass
    return None, 0

# ======================
# 3) Hàm xử lý chính
# ======================
def docx_single_table_to_df(path_docx: str) -> pd.DataFrame:
    doc = Document(path_docx)
    if not doc.tables: raise ValueError("Không tìm thấy bảng")
    table = doc.tables[0]
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    return pd.DataFrame(data[1:], columns=data[0])

def evaluate_presence(df: pd.DataFrame, text_cols: List[str], guess_values: bool, cfg: SpecsConfig) -> pd.DataFrame:
    cache = load_cache()
    out = df.copy()
    norm_cols = {norm_key(c): c for c in out.columns}
    
    col_model = next((norm_cols[k] for k in cfg.model_col_keys if k in norm_cols), None)
    col_brand = next((norm_cols[k] for k in cfg.brand_col_keys if k in norm_cols), None)
    col_mfr   = next((norm_cols[k] for k in cfg.mfr_col_keys if k in norm_cols), None)
    
    # Xác định các cột văn bản để quét
    actual_text_cols = [norm_cols[norm_key(c)] for c in text_cols if norm_key(c) in norm_cols]
    if not actual_text_cols:
        actual_text_cols = [norm_cols[norm_key(c)] for c in cfg.default_text_cols if norm_key(c) in norm_cols]

    res_m, res_b, res_f = [], [], []
    res_keyword, res_time, res_count = [], [], []

    for idx, row in out.iterrows():
        start_ts = time.time()
        m_val = str(row[col_model]).strip() if col_model and non_empty(row[col_model]) else ""
        
        # --- Trích xuất Model Nâng Cao ---
        if not m_val and guess_values:
            for c in actual_text_cols:
                text_content = str(row[c])
                # Cách 1: Tìm theo nhãn
                label_match = re.search(cfg.extract_patterns['model_labels'][0], text_content, re.I)
                if label_match:
                    m_val = label_match.group(1).strip().strip(".;:- ")
                    break
                
                # Cách 2: Quét tự do (AWC 1600)
                potential_codes = re.findall(cfg.extract_patterns['model_freeform'], text_content, re.I)
                for cand in potential_codes:
                    clean_cand = cand.strip(".;:- ")
                    cand_lower = clean_cand.lower()
                    if cand_lower in cfg.unit_blacklist: continue
                    
                    # Né các đơn vị đo dính liền (ví dụ 12.7mm)
                    pattern_with_unit = re.escape(clean_cand) + r"(?:" + "|".join(cfg.unit_blacklist) + r")\b"
                    if re.search(pattern_with_unit, text_content, re.I): continue
                    
                    m_val = clean_cand
                    break
                if m_val: break

        # --- Xử lý Hãng & Tra cứu ---
        b_val = str(row[col_brand]).strip() if col_brand and non_empty(row[col_brand]) else ""
        if b_val and norm_key(b_val) in cfg.brand_blacklist: b_val = ""
        
        found_count = 0
        keyword_used = ""
        if not b_val and m_val and guess_values:
            keyword_used = m_val # Lưu lại từ khóa mang đi tra
            if m_val in cache:
                b_val = cache[m_val]
                found_count = 1
            else:
                brand_online, count_online = fetch_brand_from_dauthau(m_val)
                if brand_online:
                    b_val = brand_online
                    cache[m_val] = brand_online
                found_count = count_online

        f_val = str(row[col_mfr]).strip() if col_mfr and non_empty(row[col_mfr]) else ""
        duration = round(time.time() - start_ts, 2)
        
        res_m.append(m_val); res_b.append(b_val); res_f.append(f_val)
        res_keyword.append(keyword_used)
        res_time.append(f"{duration}s")
        res_count.append(found_count)

    save_cache(cache)
    
    # Xuất Excel các cột kết quả
    out["Giá trị - Model"] = res_m
    out["Giá trị - Hãng"] = res_b
    out["Giá trị - Nhà sản xuất"] = res_f
    out["Tra cứu - Từ khóa sử dụng"] = res_keyword
    out["Tra cứu - Thời gian"] = res_time
    out["Tra cứu - Số dòng tìm thấy"] = res_count

    # Đánh giá
    out["Đánh giá - Model"] = ["Có" if x else "Không" for x in res_m]
    out["Đánh giá - Hãng"] = ["Có" if x else "Không" for x in res_b]
    out["Đánh giá - Nhà sản xuất"] = ["Có" if x else "Không" for x in res_f]
    
    return out

def summarize_presence(df: pd.DataFrame) -> pd.DataFrame:
    summary_data = []
    for field in ["Model", "Hãng", "Nhà sản xuất"]:
        col = f"Đánh giá - {field}"
        if col in df.columns:
            counts = df[col].value_counts()
            summary_data.append({
                "Tiêu chí": field,
                "Có": int(counts.get("Có", 0)),
                "Không": int(counts.get("Không", 0)),
                "Tổng": len(df)
            })
    return pd.DataFrame(summary_data)

def export_to_excel(evaluated: pd.DataFrame, summary: pd.DataFrame, output_path: str):
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        evaluated.to_excel(writer, sheet_name="Du_lieu", index=False)
        summary.to_excel(writer, sheet_name="Tong_hop", index=False)